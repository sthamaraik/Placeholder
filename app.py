from flask import Flask, render_template, request, send_file
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime

app = Flask(__name__)

def apply_font_style(run, font_name, size, bold=False, underline=False):
    """Apply font style to the run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # Ensure font applies to different languages
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline

def replace_text_in_paragraph(paragraph, replacements):
    """Replace text in a paragraph while preserving formatting."""
    full_text = ''.join(run.text for run in paragraph.runs)
    for key, value in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, value)

    if full_text != ''.join(run.text for run in paragraph.runs):
        for i in range(len(paragraph.runs) - 1, -1, -1):
            paragraph.runs[i].clear()
        new_run = paragraph.add_run(full_text)
        
        # Apply font styles based on placeholder
        for key in replacements:
            if key == "{{section}}" or key == "{{Date_top}}":
                apply_font_style(new_run, "Tahoma", 12, bold=True, underline=True)
            elif key == "{{Sender_Contact_No}}":
                apply_font_style(new_run, "Tahoma", 11, bold=True)
            elif key == "{{Bank_Name}}" or key == "{{Debit_freeze}}":
                apply_font_style(new_run, "Tahoma", 12, bold=True)
            elif key in ["{{FIR_CSR}}", "{{No_Year}}", "{{Date_FIR_CSR}}", "{{NCRP_No}}"]:
                apply_font_style(new_run, "Tahoma", 12)
            elif key in ["{{Mr_Mrs_Name}}", "{{Type_of_offence}}"]:
                apply_font_style(new_run, "Tahoma", 12)
            elif key in ["{{Bank_Name_letter}}", "{{ACC_No_IFSC}}", "{{Txn_ID}}", "{{Fraud_amount}}", "{{Txn_date}}"]:
                apply_font_style(new_run, "Tahoma", 11)
            elif key in ["{{Stment_date_from}}", "{{OutTxn_date}}", "{{IP_from_date}}", "{{IP_end_date}}"]:
                apply_font_style(new_run, "Tahoma", 12)

def create_table_with_placeholders(doc, replacements):
    """Create a table and replace placeholders."""
    # Create a table with one row and five columns for the placeholders
    table = doc.add_table(rows=1, cols=5)

    # Set the headers
    headers = ["Bank Name Letter", "Account No & IFSC", "Transaction ID", "Fraud Amount", "Transaction Date"]
    placeholders = ["{{Bank_Name_letter}}", "{{ACC_No_IFSC}}", "{{Txn_ID}}", "{{Fraud_amount}}", "{{Txn_date}}"]

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Add a row for the placeholder values
    row_cells = table.add_row().cells
    for i, placeholder in enumerate(placeholders):
        # Add the replacement values in the respective columns
        if placeholder in replacements:
            row_cells[i].text = replacements[placeholder]
    
    # Apply font styles to the table
    for i, placeholder in enumerate(placeholders):
        apply_font_style(row_cells[i].paragraphs[0].runs[0], "Tahoma", 11)

# Function to replace placeholders in paragraphs and tables
def replace_placeholders(doc_path, replacements):
    """Replace placeholders in the document and apply formatting."""
    doc = docx.Document(doc_path)

    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, replacements)

    # Save the modified document
    modified_doc_path = 'modified_document.docx'
    doc.save(modified_doc_path)
    return modified_doc_path

@app.route('/')
def index():
    return render_template('form.html')

@app.route('/generate', methods=['POST'])
def generate():
    """Handle form submission and generate the Word document."""
    
    # Convert date fields to dd-mm-yyyy format
    def format_date(date_string):
        try:
            return datetime.strptime(date_string, '%Y-%m-%d').strftime('%d-%m-%Y')
        except ValueError:
            return date_string  # If the date is not in a valid format, return it as-is
    
    # Replacements dictionary with formatted dates
    replacements = {
        "{{section}}": request.form['section'],
        "{{Date_top}}": format_date(request.form['Date_top']),
        "{{Sender_Contact_No}}": request.form['Sender_Contact_No'],
        "{{Bank_Name}}": request.form['Bank_Name'],
        "{{Debit_freeze}}": request.form['Debit_freeze'],
        "{{FIR_CSR}}": request.form['FIR_CSR'],
        "{{No_Year}}": request.form['No_Year'],
        "{{Date_FIR_CSR}}": format_date(request.form['Date_FIR_CSR']),
        "{{NCRP_No}}": request.form['NCRP_No'],
        "{{Mr_Mrs_Name}}": request.form['Mr_Mrs_Name'],
        "{{Type_of_offence}}": request.form['Type_of_offence'],
        "{{Bank_Name_letter}}": request.form['Bank_Name_letter'],
        "{{ACC_No_IFSC}}": request.form['ACC_No_IFSC'],
        "{{Txn_ID}}": request.form['Txn_ID'],
        "{{Fraud_amount}}": request.form['Fraud_amount'],
        "{{Txn_date}}": format_date(request.form['Txn_date']),
        "{{Stment_date_from}}": format_date(request.form['Stment_date_from']),
        "{{OutTxn_date}}": format_date(request.form['OutTxn_date']),
        "{{IP_from_date}}": format_date(request.form['IP_from_date']),
        "{{IP_end_date}}": format_date(request.form['IP_end_date'])
    }

    doc_path = 'with_placeholder.docx'
    modified_doc_path = replace_placeholders(doc_path, replacements)
    
    return send_file(modified_doc_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
